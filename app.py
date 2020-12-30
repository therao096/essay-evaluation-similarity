from flask import Flask ,jsonify,request,render_template
from flask_restful import Api , Resource
import numpy as np
import torch
import json
from sentence_transformers import SentenceTransformer
import torch
import json
import sentencepiece
from transformers import T5Tokenizer, T5ForConditionalGeneration, T5Config


app =Flask(__name__)
api = Api(app)


import gensim.downloader as gd
path = gd.load("word2vec-google-news-300", return_path=True)
model_path = path
from gensim.models.keyedvectors import KeyedVectors
w2v_model = KeyedVectors.load_word2vec_format(model_path, binary=True)
stopwords_path="D:/Accioibis_stuff/stopwords.txt"
with open(stopwords_path, 'r') as fh:
    stopwords = fh.read().split(",")


model = T5ForConditionalGeneration.from_pretrained('t5-small')
tokenizer = T5Tokenizer.from_pretrained('t5-small')

import re
import pandas as pd
from docx import Document
import warnings
warnings.filterwarnings("ignore")
def countX(lst, x): 
    return lst.count(x)



from scipy import spatial



def _cosine_sim(vecA, vecB):
        """Find the cosine similarity distance between two vectors."""
        try:
          c_sim = 1-spatial.distance.cosine(vecA,vecB)

        except:
          c_sim=0

        return c_sim




class DocSim:
    def __init__(self, w2v_model, stopwords=None):
        self.w2v_model = w2v_model
        self.stopwords = stopwords if stopwords is not None else []

    def vectorize(self, doc: str) -> np.ndarray:
        """
        Identify the vector values for each word in the given document
        :param doc:
        :return:
        """
        doc = doc.lower()
        words = [w for w in doc.split(" ") if w not in self.stopwords]
        word_vecs = []
        for word in words:
            try:
                vec = self.w2v_model[word]
                word_vecs.append(vec)
            except KeyError:
                # Ignore, if the word doesn't exist in the vocabulary
                pass

        # Assuming that document vector is the mean of all the word vectors
        # PS: There are other & better ways to do it.
        vector = np.mean(word_vecs, axis=0)
        return vector

    def _cosine_sim(self, vecA, vecB):
        """Find the cosine similarity distance between two vectors."""
        csim = np.dot(vecA, vecB) / (np.linalg.norm(vecA) * np.linalg.norm(vecB))
        if np.isnan(np.sum(csim)):
            return 0
        return csim

    def calculate_similarity(self, source_doc, target_docs=None, threshold=0):
        """Calculates & returns similarity scores between given source document & all
        the target documents."""
        if not target_docs:
            return []

        if isinstance(target_docs, str):
            target_docs = [target_docs]

        source_vec = self.vectorize(source_doc)
        results = []
        for doc in target_docs:
            target_vec = self.vectorize(doc)
            sim_score = self._cosine_sim(source_vec, target_vec)
            if sim_score > threshold:
                results.append( sim_score)
            # Sort results by score in desc order
            #results.sort(key=lambda k: k["score"], reverse=True)

        return results




ds = DocSim(w2v_model,stopwords=stopwords)

@app.route('/')
def my_form():
    return render_template("input_form.html")

@app.route('/', methods=['POST'])
def Get_grades_org():
    paragraph = request.form["essay"]
    prob = request.form["question"]


    paar=paragraph.replace("\n\n","\n")
  
    para=paar.split("\n")
  
    while '' in para:para.remove('')
      #para.remove('')
      #para.remove(' ')
    #except Exception:
      #abc=len(para)
    abc=len(para)
    #print(para)
    #print(abc)
    #print(prob)
    lines=[]
    lines_res=[]
    for j in range(0,abc):
      z=para[j]
      y=z.split(".")
      ee=y[0]
    #xe=y[1].join(y[:])
      lines.append(ee)
      for e in y:
        s=1
        le=len(y)
        att=''.join(y[1:le])
      lines_res.append(att)
    #print(lines_res)

    summarys=[]
    for i in para:
        text=i
        preprocess_text = text.strip().replace("\n","")
        t5_prepared_Text = "summarize: "+preprocess_text
        #print ("original text preprocessed: \n", preprocess_text)

        tokenized_text = tokenizer.encode(t5_prepared_Text, return_tensors="pt")


        # summmarize 
        summary_ids = model.generate(tokenized_text,
                                        num_beams=4,
                                        no_repeat_ngram_size=2,
                                        min_length=30,
                                        max_length=100,
                                        early_stopping=True)

        output = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
        summarys.append(output)

    texts=prob
    preprocess_texts = texts.strip().replace("\n","")
    t5_prepared_Texts = "summarize: "+preprocess_texts
        #print ("original text preprocessed: \n", preprocess_text)

    tokenized_texts = tokenizer.encode(t5_prepared_Texts, return_tensors="pt")


        # summmarize 
    summary_idss = model.generate(tokenized_texts,
                                      num_beams=4,
                                      no_repeat_ngram_size=2,
                                      min_length=30,
                                      max_length=100,
                                      early_stopping=True)

    outputs = tokenizer.decode(summary_idss[0], skip_special_tokens=True)
    #print(outputs)
   #first line of para1
    texts_first=lines[0]
    preprocess_texts_first = texts_first.strip().replace("\n","")
    t5_prepared_Texts_first = "summarize: "+preprocess_texts_first
        #print ("original text preprocessed: \n", preprocess_text)

    tokenized_texts_first = tokenizer.encode(t5_prepared_Texts_first, return_tensors="pt")


        # summmarize 
    summary_idsss = model.generate(tokenized_texts_first,
                                      num_beams=4,
                                      no_repeat_ngram_size=2,
                                      min_length=30,
                                      max_length=100,
                                      early_stopping=True)

    outputs_firstline = tokenizer.decode(summary_idsss[0], skip_special_tokens=True)
    #print(outputs_firstline)
  #remaining line of para1
    texts_rest=lines_res[0]
    preprocess_texts_rest = texts_rest.strip().replace("\n","")
    t5_prepared_Texts_rest = "summarize: "+preprocess_texts_rest
        #print ("original text preprocessed: \n", preprocess_text)

    tokenized_texts_rest = tokenizer.encode(t5_prepared_Texts_rest, return_tensors="pt")


        # summmarize 
    summary_idssss = model.generate(tokenized_texts_rest,
                                      num_beams=4,
                                      no_repeat_ngram_size=2,
                                      min_length=30,
                                      max_length=100,
                                      early_stopping=True)

    outputs_rest = tokenizer.decode(summary_idssss[0], skip_special_tokens=True)
    #print(lines_res[0])
    #print(outputs_rest)




    lines_sum=[]
    lines_res_sum=[]
    for c in range(0,abc):
      w=summarys[c]
      yy=w.split(".")
      eee=yy[0]
      #xe=y[1].join(y[:])
      lines_sum.append(eee)
      for es in yy:
        ss=1
        lee=len(yy)
        attt=''.join(yy[1:lee])
      lines_res_sum.append(attt)
    #print(summarys[0])
   # print(lines_sum)
    #print(lines_res_sum)
    #while '' in lines_sum:lines_sum.remove('')
    #while '' in lines_res_sum:lines_res_sum.remove('')

    sim_scores1=ds.calculate_similarity(outputs, outputs_firstline)
    sim_scores2=ds.calculate_similarity(outputs, outputs_rest)
    try:
      sim_scores3=ds.calculate_similarity(outputs, summarys[1])
    except Exception as e:
      sim_scores3=0
    try:
      sim_scores4=ds.calculate_similarity(outputs, summarys[2])
    except Exception as e:
      sim_scores4=0
    try:
      sim_scores5=ds.calculate_similarity(summarys[1], summarys[2])
    except Exception as e:
      sim_scores5=0
  ##linking words
    document1 = Document('D:/Accioibis_stuff/link.docx')
    link=[]
    hello=['','\xa0']
    for paras in document1.paragraphs:
        if paras.text not in hello:
        #if paras.text != '':
          link.append(paras.text)


    length=len(para)
    words=[]
    parass=[]
    for ss in range(0,length):
        essay=para[ss]
        for j in link:
            if re.search(j,essay,re.IGNORECASE):
                words.append(j)
                parass.append(ss)
      
      #valu=Counter(parass)

    coo=[]
    for co in range(0,length):
      coo.append(countX(parass,co))
    a_list=zip(words,parass)
    final_list=list(a_list)
    df=pd.DataFrame()
    df['Words']=words
    df['Para_number']=parass
    tot_num=df.Para_number.count()

    #results1=[]
    #results2=[]
    try:
      results1=[]
      for me in coo:
        #res_avg=(me/tot_num)
        res1=(me/tot_num)**2
        results1.append(res1)
        #results2.append(res_avg)
      total=sum(results1)
      total_avg= tot_num/(length)
    except Exception as e:
      total=0
      total_avg=0
  ##reference words
    document2 = Document('D:/Djangoprojs/reference.docx')
    reference=[]
    hello=['','\xa0']
    for paras in document2.paragraphs:
        if paras.text not in hello:
        #if paras.text != '':
          reference.append(paras.text)

    length=len(para)
    wordsr=[]
    parassr=[]
    for ssr in range(0,length):
        essayr=para[ssr]
        for jr in reference:
            if re.search(jr,essayr,re.IGNORECASE):
                wordsr.append(jr)
                parassr.append(ssr)
      
      #valu=Counter(parass)
    coor=[]
    for cor in range(0,length):
      coor.append(countX(parassr,cor))
    a_listr=zip(wordsr,parassr)
    final_listr=list(a_listr)
    dfr=pd.DataFrame()
    dfr['Words']=wordsr
    dfr['Para_number']=parassr
    tot_numr=dfr.Para_number.count()

    try:
      results2=[]
      for mer in coor:
        res2=(mer/tot_numr)**2
        results2.append(res2)
      totalr=sum(results2)
      avg_ref= tot_numr/(length)
      countr=dfr.groupby(['Para_number']).size().reset_index(name='counts')
    except Exception as e:
      totalr=0
      countr=0

  #conclusion 
    document3 = Document('D:/Djangoprojs/conclusion.docx')
    concl=[]
    hello=['','\xa0']
    for paras in document3.paragraphs:
        if paras.text not in hello:
        #if paras.text != '':
          concl.append(paras.text)
    try:
      sim_scores6 = ds.calculate_similarity(summarys[abc-1],outputs_rest)
    except Exception as e:
      sim_scores6=0
    #length=len(para)
    wordsc=[]
    parassc=[]
    essayc=para[-1]
    #no=essayc.split()
    #jj=len(no)
    try:
      for jc in concl:
          if re.search(jc,essayc,re.IGNORECASE):
                  wordsc.append(jc)
                    #parassc.append(ssc)
      #print(wordsc)
      no_concl=len(wordsc)
    except Exception as e:
      no_concl=0

    #print(para[0])
    #print(lines[0])
    #print(lines_res[0])
    #print(outputs_firstline)
    #print(outputs_rest)
    #return exe
    results=[abc,str(sim_scores1),str(sim_scores2),str(sim_scores3),str(sim_scores4),str(sim_scores5),str(no_concl),str(sim_scores6),str(total_avg),str(total),str(avg_ref),str(totalr)]
    #results=[abc,a1,b1,c1,d1,e1,str(no_concl),f1,str(total_avg),str(total),str(avg_ref),str(totalr)]
    output = f'len = {str(abc)} , sim_scores1 = {str(sim_scores1)} , sim_scores2 = {str(sim_scores2)} , sim_scores3 = {str(sim_scores3)} , sim_scores4 = {str(sim_scores4)} , sim_scores5 = {str(sim_scores5)} , sim_scores6 = {str(sim_scores6)} , total_avg , total , avg_ref , totalr = {str(total_avg)} , {str(total)} , {str(avg_ref)} , {str(totalr)} '
    return output


if __name__ == "__main__":
    app.run()
